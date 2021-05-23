# TODO Run command to install Python Packages --> pip3 install -r requirements.txt
# IMPORT python-docx
from docx import Document
from docx.shared import Inches  # RESIZING PICS
import pyttsx3  # TEXT-TO-SPEECH

# TODO Add validations for user inputs

# TEXT-TO-SPEECH
pyttsx3.speak('This is an easy solution to build your CV')


# TAKE INPUT TO SPEAK
def speak(text):
    pyttsx3.speak(text)


# CREATING DOCUMENT OBJECT
document = Document()

# ADDING PICTURE
document.add_picture('cv_pic.jpg', width=Inches(1.5))

# TAKING INPUT FROM USER
name = input('What is your name? ')
speak(f'Hello {name}, How are you today?')
phone_number = input('What is your phone number? ')
email = input('What is your email-id? ')

# ADD TEXT TO DOC
document.add_paragraph(f"{name} | {phone_number} | {email}")

# HEADING
document.add_heading("About Me")
about_me = input("Tell us about yourself: ")
document.add_paragraph(about_me)

# WORK EXPERIENCE
document.add_heading("Work Experience")
p = document.add_paragraph()
company = input("Enter company: ")
from_date = input('From date: ')
to_date = input('To date: ')
# ADD DATA TO EXISTING PARAGRAPHS
p.add_run(company + ' ').bold = True  # BOLD
p.add_run(from_date + ' - ' + to_date + '\n').italic = True  # ITALIC

experience_details = input(f'Describe your experience at {company}: ')
p.add_run(experience_details)

# MORE EXPERIENCES
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input("Enter company: ")
        from_date = input('From date: ')
        to_date = input('To date: ')
        # ADD DATA TO EXISTING PARAGRAPHS
        p.add_run(company + ' ').bold = True  # BOLD
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True  # ITALIC

        experience_details = input(f'Describe your experience at {company}: ')
        p.add_run(experience_details)
    else:
        break

# SKILLS
document.add_heading('Skills')
skill = input('Enter your skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'  # LIST AS BULLET POINTS
# MORE SKILLS
while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter your skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'  # LIST AS BULLET POINTS
    else:
        break

# FOOTER
section = document.sections[0]  # GRAB A SECTION
footer = section.footer  # GRAB A FOOTER
p = footer.paragraphs[0]  # GRAB A PARAGRAPH
p.text = 'CV generated by Aanif while learning Python'

# NAMING AND SAVING INTO FILE
document.save("cv.docx")
